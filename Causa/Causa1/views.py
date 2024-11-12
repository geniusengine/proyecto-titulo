from django.shortcuts import render, redirect,get_object_or_404
from django.contrib import messages
from django.views.decorators.csrf import csrf_protect
from django.contrib.auth import authenticate, login
from django.db import connection, transaction
import logging
from django.contrib.auth.models import User
from django.db import IntegrityError
from .models import Notificacion
from .models import Usuario
import hashlib 
import bcrypt
from django.http import HttpResponse
from docx import Document
from datetime import datetime
from .models import Estampado
from .forms import DemandaForm
from .models import Demanda
from django.urls import reverse
from django.contrib.auth import logout
from django.utils import timezone
from datetime import timedelta
from .models import AUD_notificacion

def logout_view(request):
    logout(request)
    messages.success(request, "Sesión cerrada correctamente.")
    return redirect('main_menu')




# menuprincipal
@csrf_protect
def main_menu(request):
    return render(request, 'Causa1/main_menu.html')

@csrf_protect
def login_view(request):
    if request.method == 'POST':
        # Lógica de inicio de sesión aquí (similar a `open_login`)
        return render(request, 'Causa1/login.html')  # página de login

@csrf_protect
def register_view(request):
    if request.method == 'POST':
        # Usuario y contraseña designados
        usuario_designado = "admin"
        contrasena_designada = "admin"
        
        # Obtiene usuario y contraseña desde el formulario
        usuario = request.POST.get('usuario')
        contrasena = request.POST.get('contrasena')
        
        # Verifica las credenciales
        if usuario == usuario_designado and contrasena == contrasena_designada:
            return render(request, 'Causa1/register.html')  # página de registro
        else:
            messages.error(request, "Usuario y/o contraseña incorrectos.")
            return redirect('main_menu')
    return redirect('main_menu')

# Configuración del sistema de login
logging.basicConfig(filename='registro.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')

        if username and password:
            try:
                # Busca al usuario por su nombre de usuario
                user = Usuario.objects.get(username=username)
                
                # Verificar si la contraseña está hasheada con bcrypt o MD5
                if len(user.password) == 32:  # Probable hash MD5 (32 caracteres)
                    # Generar el hash MD5 de la contraseña ingresada
                    hashed_password_md5 = hashlib.md5(password.encode()).hexdigest()
                    
                    # Verificar si el hash MD5 coincide
                    if hashed_password_md5 == user.password:
                        # Hashear la contraseña con bcrypt y actualizar la base de datos
                        hashed_password_bcrypt = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()
                        user.password = hashed_password_bcrypt
                        user.save()
                        messages.info(request, "Tu contraseña ha sido actualizada para mayor seguridad.")

                        # Almacenar el ID del usuario en la sesión
                        request.session['user_id'] = user.id
                        messages.success(request, "Inicio de sesión exitoso.")
                        return redirect('dashboard')
                    else:
                        messages.error(request, "Usuario o contraseña incorrectos.")
                else:
                    # Si ya está en bcrypt, verificar directamente
                    if bcrypt.checkpw(password.encode(), user.password.encode()):
                        # Almacenar el ID del usuario en la sesión
                        request.session['user_id'] = user.id
                        messages.success(request, "Inicio de sesión exitoso.")
                        return redirect('dashboard')
                    else:
                        messages.error(request, "Usuario o contraseña incorrectos.")
            except Usuario.DoesNotExist:
                messages.error(request, "Usuario o contraseña incorrectos.")
        else:
            messages.error(request, "Por favor, ingresa ambos campos.")

    return render(request, 'Causa1/login.html')


# Registro

def register_view(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        apellido = request.POST.get('apellido')
        username = request.POST.get('username')
        password = request.POST.get('password')
        
        # Verificar campos vacíos
        if not name or not apellido or not username or not password:
            messages.error(request, "Por favor, complete todos los campos.")
            return render(request, 'Causa1/register.html')
        
        # Hashear la contraseña usando bcrypt
        hashed_password = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

        # Crear el usuario usando el modelo Usuario
        try:
            user = Usuario.objects.create(
                username=username,
                nombreusuario=name,
                apellidousuario=apellido,
                password=hashed_password
            )
            messages.success(request, "Usuario registrado correctamente.")
            return redirect('login')
        except IntegrityError:
            messages.error(request, "El nombre de usuario ya existe.")
            return render(request, 'Causa1/register.html')

    return render(request, 'Causa1/register.html')


#Dashboar 
def dashboard(request):
    # Verificar si el usuario ha iniciado sesión
    if 'user_id' not in request.session:
        messages.error(request, "Debes iniciar sesión primero.")
        return redirect('login')
    
    # Eliminar demandas en verde con más de 5 minutos
    tiempo_limite = timezone.now() - timedelta(minutes=5)
    Notificacion.objects.filter(estadoNoti=True, estadoCausa=True, fechaNotificacion__lt=tiempo_limite).delete()
    
    # Cargar los datos de Notificacion si el usuario está autenticado
    causas = Notificacion.objects.all()
    return render(request, 'Causa1/dashboard.html', {'causas': causas})

def notificar(request, causa_id):
    causa = get_object_or_404(Notificacion, id=causa_id)
    if request.method == "POST":
        if causa.estadoNoti == 0:  # Solo cambia si es 0
            causa.estadoNoti = 1  # Cambia el valor a 1
            causa.save()
            messages.success(request, "Causa notificada correctamente.")
        else:
            messages.warning(request, "Esta causa ya ha sido notificada.")
    return redirect('dashboard')

def estampar(request, causa_id):
    causa = get_object_or_404(Notificacion, id=causa_id)
    
    if request.method == "POST":
        tipo_estampado = request.POST.get("tipo_estampado")
        
        # Verificamos si la causa puede ser estampada
        if causa.estadoNoti == 1:
            if causa.estadoCausa == 0:
                causa.estadoCausa = 1
                causa.save()
                messages.success(request, "Causa estampada correctamente.")
                
                # Redirigir a la función de descarga con los parámetros necesarios
                return redirect('descargar_documento', estampado_id=causa.id, tipo_estampado=tipo_estampado)
            else:
                messages.warning(request, "Esta causa ya ha sido estampada.")
        else:
            messages.warning(request, "La causa debe ser notificada antes de estampar.")
    
    return redirect('dashboard')

def descargar_documento(request, estampado_id, tipo_estampado):
    notificacion = get_object_or_404(Notificacion, id=estampado_id)
    doc = Document()
    
    # Fecha y hora actuales
    now = datetime.now()
    fecha_actual = now.strftime("%d/%m/%Y")
    hora_actual = now.strftime("%H:%M")

    # Encabezado común
    encabezado = f"{notificacion.nombTribunal}\nCAUSA ROL= {notificacion.numjui}\nCARÁTULA {notificacion.demandante} / {notificacion.demandado}"
    doc.add_paragraph(encabezado)

    # Contenido específico según el tipo de estampado
    if tipo_estampado == "negativa52":
        contenido = f"BÚSQUEDA NEGATIVA: Certifico haber buscado al(la) demandado(a) {notificacion.demandado}, con domicilio en {notificacion.domicilio} {notificacion.comuna} especialmente el día {fecha_actual}, siendo las {hora_actual} horas..."
    elif tipo_estampado == "positivaP":
        contenido = f"BÚSQUEDA POSITIVA: En {fecha_actual}, siendo las {hora_actual}, en {notificacion.domicilio}..."
    elif tipo_estampado == "busquedaN":
        contenido = f"BÚSQUEDA Y NOTIFICACIÓN: a {fecha_actual}, en {notificacion.domicilio}..."

    # Firma y cierre
    doc.add_paragraph(contenido)
    doc.add_paragraph(f"Drs. {notificacion.arancel}.-")

    # Preparar el archivo para la descarga
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{notificacion.numjui}_{tipo_estampado}.docx"'
    doc.save(response)
    return response

def estampado(request, estampado_id):
    # Obtén la notificación específica basada en el ID
    notificacion = get_object_or_404(Notificacion, id=estampado_id)
    return render(request, 'Causa1/estampado.html', {'notificacion': notificacion})


#insertar datos

from django.shortcuts import render, redirect
from django.db import connection  # Para conexión directa con la base de datos
from .forms import DemandaForm

def crear_demanda(request):
    if request.method == "POST":
        form = DemandaForm(request.POST)
        if form.is_valid():
            # Guardar en la tabla 'demanda'
            with connection.cursor() as cursor:
                cursor.execute("""
                    INSERT INTO demanda (numjui, nombTribunal, demandante, demandado, repre, mandante, domicilio, comuna, encargo, soli, arancel, actu)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, [
                    form.cleaned_data['numjui'],
                    form.cleaned_data['nombTribunal'],
                    form.cleaned_data['demandante'],
                    form.cleaned_data['demandado'],
                    form.cleaned_data['repre'],
                    form.cleaned_data['mandante'],
                    form.cleaned_data['domicilio'],
                    form.cleaned_data['comuna'],
                    form.cleaned_data['encargo'],
                    form.cleaned_data['soli'],
                    form.cleaned_data['arancel'],
                    form.cleaned_data['actu']
                ])
            return redirect('dashboard')

    else:
        form = DemandaForm()

    # Extrae las opciones de los campos 'nombTribunal' y 'actu' para el template
    tribunal_choices = form.fields['nombTribunal'].choices
    actu_choices = form.fields['actu'].choices
    arancel_choices = form.ARANCELES_CHOICES

    return render(request, 'Causa1/crear_demanda.html', {
        'tribunal_choices': tribunal_choices,
        'actu_choices': actu_choices,
        'arancel_choices': arancel_choices,
        'form': form
    })




#historico

def dashboard_historico(request):
    # Obtén todos los registros de la tabla AUD_notificacion
    historico = AUD_notificacion.objects.all()
    return render(request, 'Causa1/dashboard_historico.html', {'historico': historico})