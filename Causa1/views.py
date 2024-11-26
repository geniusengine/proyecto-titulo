from django.shortcuts import render, redirect,get_object_or_404
from django.contrib import messages
from django.views.decorators.csrf import csrf_protect
from django.contrib.auth import authenticate, login
from django.db import connection, transaction
import logging
from django.contrib.auth.models import User
from django.db import IntegrityError
from .models import Notificacion
from .models import Usuario
import hashlib 
import bcrypt
from django.http import HttpResponse
from docx import Document
from datetime import datetime

from .forms import DemandaForm
from .models import Demanda
from django.urls import reverse
from django.contrib.auth import logout
from django.utils import timezone
from datetime import timedelta
from .models import AUD_notificacion
from django.contrib.messages import get_messages

def logout_view(request):
    logout(request)
    messages.success(request, "Sesión cerrada correctamente.")
    # Recupera los mensajes y asegura que desaparezcan después de mostrarse
    almacen_mensajes = get_messages(request)
    return redirect('main_menu')




# menuprincipal
@csrf_protect
def main_menu(request):
    # Recupera los mensajes y asegura que desaparezcan después de mostrarse
    almacen_mensajes = get_messages(request)
    return render(request, 'Causa1/main_menu.html',{'messages': almacen_mensajes})

@csrf_protect
def login_view(request):
    # Recupera los mensajes y asegura que desaparezcan después de mostrarse
    almacen_mensajes = get_messages(request)
    if request.method == 'POST':
        # Lógica de inicio de sesión aquí (similar a `open_login`)
        return render(request, 'Causa1/login.html',{'messages': almacen_mensajes})  # página de login

@csrf_protect
def register_view(request):
    # Recupera los mensajes y asegura que desaparezcan después de mostrarse
    almacen_mensajes = get_messages(request)
    if request.method == 'POST':
        # Usuario y contraseña designados
        usuario_designado = "admin"
        contrasena_designada = "admin"
        
        # Obtiene usuario y contraseña desde el formulario
        usuario = request.POST.get('usuario')
        contrasena = request.POST.get('contrasena')
        
        # Verifica las credenciales
        if usuario == usuario_designado and contrasena == contrasena_designada:
            return render(request, 'Causa1/register.html',{'messages': almacen_mensajes})  # página de registro
        else:
            messages.error(request, "Usuario y/o contraseña incorrectos.")
            return redirect('main_menu')
    return redirect('main_menu')

# Configuración del sistema de login
logging.basicConfig(filename='registro.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def login_view(request):
    almacen_mensajes = get_messages(request)
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        # Recupera los mensajes y asegura que desaparezcan después de mostrarse
      

        if username and password:
            try:
                # Busca al usuario por su nombre de usuario
                user = Usuario.objects.get(username=username)
                
                # Verificar si la contraseña está hasheada con bcrypt o MD5
                if len(user.password) == 32:  # Probable hash MD5 (32 caracteres)
                    # Generar el hash MD5 de la contraseña ingresada
                    hashed_password_md5 = hashlib.md5(password.encode()).hexdigest()
                    
                    # Verificar si el hash MD5 coincide
                    if hashed_password_md5 == user.password:
                        # Hashear la contraseña con bcrypt y actualizar la base de datos
                        hashed_password_bcrypt = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()
                        user.password = hashed_password_bcrypt
                        user.save()
                        messages.info(request, "Tu contraseña ha sido actualizada para mayor seguridad.")

                        # Almacenar el ID del usuario en la sesión
                        request.session['user_id'] = user.id
                        messages.success(request, "Inicio de sesión exitoso.")
                        return redirect('dashboard')
                    else:
                        messages.error(request, "Usuario o contraseña incorrectos.")
                else:
                    # Si ya está en bcrypt, verificar directamente
                    if bcrypt.checkpw(password.encode(), user.password.encode()):
                        # Almacenar el ID del usuario en la sesión
                        request.session['user_id'] = user.id
                        messages.success(request, "Inicio de sesión exitoso.")
                        return redirect('dashboard')
                    else:
                        messages.error(request, "Usuario o contraseña incorrectos.")
            except Usuario.DoesNotExist:
                messages.error(request, "Usuario o contraseña incorrectos.")
        else:
            messages.error(request, "Por favor, ingresa ambos campos.")

    return render(request, 'Causa1/login.html',{'messages': almacen_mensajes})


# Registro

def register_view(request):
    # Recupera los mensajes y asegura que desaparezcan después de mostrarse
    almacen_mensajes = get_messages(request)
    if request.method == 'POST':
        name = request.POST.get('name')
        apellido = request.POST.get('apellido')
        username = request.POST.get('username')
        password = request.POST.get('password')
        
        # Verificar campos vacíos
        if not name or not apellido or not username or not password:
            messages.error(request, "Por favor, complete todos los campos.")
            return render(request, 'Causa1/register.html',)
        
        # Hashear la contraseña usando bcrypt
        hashed_password = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

        # Crear el usuario usando el modelo Usuario
        try:
            user = Usuario.objects.create(
                username=username,
                nombreusuario=name,
                apellidousuario=apellido,
                password=hashed_password
            )
            messages.success(request, "Usuario registrado correctamente.")
            return redirect('login')
        except IntegrityError:
            messages.error(request, "El nombre de usuario ya existe.")
            return render(request, 'Causa1/register.html')

    return render(request, 'Causa1/register.html',{'messages': almacen_mensajes})


#Dashboar 



from django.utils.timezone import now
from datetime import timedelta

def dashboard(request):
    """
    Vista principal del dashboard que muestra las notificaciones.
    """
    almacen_mensajes = get_messages(request)
    # Verificar si el usuario ha iniciado sesión
    if 'user_id' not in request.session:
        messages.error(request, "Debes iniciar sesión primero.")
        return redirect('login')

    # Obtener el usuario actual
    try:
        user_id = request.session['user_id']
        usuario = Usuario.objects.get(id=user_id)
    except Usuario.DoesNotExist:
        messages.error(request, "Usuario no encontrado.")
        return redirect('login')

    # Filtrar causas activas (estadoNoti y estadoCausa en 1)
    tiempo_limite = now() - timedelta(minutes=5)
    causas_verdes = Notificacion.objects.filter(
        estadoNoti=1, estadoCausa=1, fechaNotificacion__lte=tiempo_limite
    )

    # Mover causas a `AUD_notificacion`
    for causa in causas_verdes:
        AUD_notificacion.objects.create(
            fechaNotificacion=causa.fechaNotificacion,
            numjui=causa.numjui,
            nombTribunal=causa.nombTribunal,
            demandante=causa.demandante,
            demandado=causa.demandado,
            repre=causa.repre,
            mandante=causa.mandante,
            domicilio=causa.domicilio,
            comuna=causa.comuna,
            encargo=causa.encargo,
            soli=causa.soli,
            arancel=causa.arancel,
            arancel_nombre=causa.arancel_nombre,
            estadoNoti=causa.estadoNoti,
            estadoCausa=causa.estadoCausa,
            actu=causa.actu
        )
        # Eliminar causa de `Notificacion`
        causa.delete()

    # Cargar los datos restantes para el dashboard
    causas = Notificacion.objects.all().values(
        'id',
        'fechaNotificacion',
        'numjui',
        'nombTribunal',
        'demandante',
        'demandado',
        'repre',
        'mandante',
        'domicilio',
        'comuna',
        'encargo',
        'soli',
        'arancel',
        'arancel_nombre',
        'actu',
        'estadoNoti',
        'estadoCausa'
    )

    return render(request, 'Causa1/dashboard.html', {
        'causas': causas,
        'nombreusuario': usuario.nombreusuario,
        'apellidousuario': usuario.apellidousuario,
        'messages': almacen_mensajes
    })

def notificar(request, causa_id):
    """
    Cambiar el estado de notificación de una causa específica.
    """
    causa = get_object_or_404(Notificacion, id=causa_id)

    if request.method == "POST":
        if causa.estadoNoti == 0:  # Solo permite cambiar si el estado es 0 (no notificado)
            causa.estadoNoti = 1  # Cambiar el estado a 1 (notificado)
            causa.save()
            messages.success(request, "Causa notificada correctamente.")
        else:
            messages.warning(request, "Esta causa ya ha sido notificada.")

    return redirect('dashboard')


def estampar(request, causa_id):
    """
    Cambiar el estado de estampado de una causa específica y generar un documento si aplica.
    """
    causa = get_object_or_404(Notificacion, id=causa_id)

    if request.method == "POST":
        tipo_estampado = request.POST.get("tipo_estampado")

        # Verificar si la causa puede ser estampada
        if causa.estadoNoti == 1:  # Solo se puede estampar si la causa ya está notificada
            if causa.estadoCausa == 0:  # Solo permite estampar si aún no está estampada
                causa.estadoCausa = 1  # Cambiar el estado a estampado
                causa.fechaNotificacion = timezone.now()  # Actualizar la fecha de notificación
                causa.save()
                messages.success(request, "Causa estampada correctamente.")

                # Redirigir a la función de descarga del documento
                return redirect('descargar_documento', estampado_id=causa.id, tipo_estampado=tipo_estampado)
            else:
                messages.warning(request, "Esta causa ya ha sido estampada.")
        else:
            messages.warning(request, "La causa debe ser notificada antes de estampar.")

    return redirect('dashboard')







def descargar_documento(request, estampado_id, tipo_estampado):
    # Obtener la notificación específica
    notificacion = get_object_or_404(Notificacion, id=estampado_id)
    doc = Document()
    
    # Fecha y hora actuales
    now = datetime.now()
    fecha_actual = now.strftime("%d/%m/%Y")
    hora_actual = now.strftime("%H:%M")

    # Encabezado común
    encabezado = (
        "Alejandra Muñoz Orellana\n"
        "Receptora Judicial – La Serena\n"
        "receptoralejandramunoz@gmail.com\n"
        "Av. Del Mar, N° 5.700, of. N° 47  La Serena.\n"
        "+56 952178958\n\n"
        f"{notificacion.nombTribunal}\n"
        f"ROL-RIT: {notificacion.numjui}\n"
        f"{notificacion.demandante} CON {notificacion.demandado}\n"
    )
    doc.add_paragraph(encabezado)

    # Contenido específico según el tipo de estampado
    if tipo_estampado == "negativa52":
        contenido = (
            f"BÚSQUEDA NEGATIVA: Certifico haber buscado al(la) demandado(a) "
            f"{notificacion.demandado}, con domicilio en {notificacion.domicilio}, {notificacion.comuna}, "
            f"especialmente el día {fecha_actual}, siendo las {hora_actual} horas, a fin de notificarle la resolución "
            f"de fecha . Diligencia que no se llevó a efecto por cuanto el(la) demandado(a) "
            f"no fue habido(a). DOY FE."
        )
    elif tipo_estampado == "positivaP":
        contenido = (
            f"BÚSQUEDA POSITIVA: a {fecha_actual}, siendo las {hora_actual} horas, en su domicilio ubicado en "
            f"{notificacion.domicilio}, {notificacion.comuna}, busqué a {notificacion.demandado}, "
            f"a fin de notificarle la demanda íntegra y su respectivo proveído. Diligencia que no se llevó a efecto "
            f"por no ser habido en dicho domicilio, en ese momento. DOY FE."
        )
    elif tipo_estampado == "busquedaN":
        contenido = (
            f"BÚSQUEDA Y NOTIFICACIÓN: a {fecha_actual}, siendo las {hora_actual} horas, en su domicilio ubicado en "
            f"{notificacion.domicilio}, {notificacion.comuna}, busqué a {notificacion.demandado}, "
            f"a fin de notificarle la resolución de fecha {notificacion.fecha_resolucion}, junto al escrito que antecede. "
            f"Diligencia que no se llevó a efecto por no ser habido en dicho domicilio, en ese momento. DOY FE."
        )

    # Agregar contenido y firma
    doc.add_paragraph(contenido)
    doc.add_paragraph(f"Drs. {notificacion.arancel_nombre} - ${notificacion.arancel}.-")

    # Preparar el archivo para la descarga
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{notificacion.numjui}_{tipo_estampado}.docx"'
    doc.save(response)
    return response




def crear_demanda(request):
    if request.method == "POST":
        form = DemandaForm(request.POST)
        if form.is_valid():
            # Obtén los valores del formulario directamente
            arancel_nombre = request.POST.get('arancel_nombre')
            arancel_valor = request.POST.get('arancel')

            # Verifica si los valores son válidos
            try:
                arancel_valor = int(arancel_valor)  # Convertir el valor a entero
            except ValueError:
                messages.error(request, "El valor del arancel es inválido.")
                return redirect('crear_demanda')

            # Crear la instancia de Demanda
            try:
                demanda = Demanda.objects.create(
                    numjui=form.cleaned_data['numjui'],
                    nombTribunal=form.cleaned_data['nombTribunal'],
                    demandante=form.cleaned_data['demandante'],
                    demandado=form.cleaned_data['demandado'],
                    repre=form.cleaned_data['repre'],
                    mandante=form.cleaned_data['mandante'],
                    domicilio=form.cleaned_data['domicilio'],
                    comuna=form.cleaned_data['comuna'],
                    encargo=form.cleaned_data['encargo'],
                    soli=form.cleaned_data['soli'],
                    arancel=arancel_valor,
                    arancel_nombre=arancel_nombre,
                    actu=form.cleaned_data['actu']
                )

                # Crear la instancia de Notificacion con los mismos datos
                Notificacion.objects.create(
                    fechaNotificacion=timezone.now(),
                    numjui=demanda.numjui,
                    nombTribunal=demanda.nombTribunal,
                    demandante=demanda.demandante,
                    demandado=demanda.demandado,
                    repre=demanda.repre,
                    mandante=demanda.mandante,
                    domicilio=demanda.domicilio,
                    comuna=demanda.comuna,
                    encargo=demanda.encargo,
                    soli=demanda.soli,
                    arancel=demanda.arancel,
                    arancel_nombre=demanda.arancel_nombre,
                    actu=demanda.actu,
                    estadoNoti=0,  # Inicializa el estado como no notificado
                    estadoCausa=0  # Inicializa el estado como pendiente
                )

                messages.success(request, "Demanda y Notificación creadas exitosamente.")
                return redirect('lista_demandas')
            except Exception as e:
                messages.error(request, f"Error al crear la demanda: {e}")
                return redirect('crear_demanda')
    else:
        form = DemandaForm()

    # Pasar opciones adicionales a la plantilla
    tribunal_choices = form.fields['nombTribunal'].choices
    actu_choices = form.fields['actu'].choices
    arancel_choices = form.ARANCELES_CHOICES

    return render(request, 'Causa1/crear_demanda.html', {
        'form': form,
        'tribunal_choices': tribunal_choices,
        'actu_choices': actu_choices,
        'arancel_choices': arancel_choices,
    })







#historico

def dashboard_historico(request):
    # Obtén todos los registros de la tabla AUD_notificacion
    # Recupera los mensajes y asegura que desaparezcan después de mostrarse
    almacen_mensajes = get_messages(request)
    historico = AUD_notificacion.objects.all()
    return render(request, 'Causa1/dashboard_historico.html', {'historico': historico,
                                                               'messages': almacen_mensajes})